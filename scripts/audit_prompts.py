import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add src to path
sys.path.append(str(Path(__file__).resolve().parent.parent))
from src.prompts import PromptGenerator

def create_audit_report(output_file: str = "Prompt_Audit_Report.docx"):
    doc = Document()
    
    # Title
    title = doc.add_heading('Prompt Strategy Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('This document contains the exact prompts generated for each strategy in Few-shot mode, used to audit the correctness of system instructions, examples, and user input formatting.')
    
    # Mock Data
    mock_title = "Urban renewal of public housing in Shenzhen: A case study"
    mock_abstract = "This study examines the redevelopment of Futian district in Shenzhen, China. It analyzes the spatial impact of renewal projects."
    
    # Strategies to audit
    strategies = ["single", "stepwise", "cot", "reflection"]
    # Note: stepwise_long shares the same prompts as stepwise, just different memory management.
    # So checking stepwise covers the prompt content for stepwise_long too.
    
    prompt_gen = PromptGenerator(shot_mode="few")
    
    for strategy in strategies:
        doc.add_page_break()
        heading = doc.add_heading(f"Strategy: {strategy.upper()} (Few-shot)", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        
        # 1. System Prompt
        doc.add_heading("1. System Prompt (including Examples)", level=2)
        
        if strategy == "single":
            sys_prompt = prompt_gen.get_single_system_prompt()
        elif strategy == "stepwise":
            sys_prompt = prompt_gen.get_step_system_prompt()
        elif strategy == "cot":
            sys_prompt = prompt_gen.get_cot_system_prompt()
        elif strategy == "reflection":
            sys_prompt = prompt_gen.get_reflection_system_prompt()
        else:
            sys_prompt = "Unknown Strategy"
            
        p = doc.add_paragraph(sys_prompt)
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(9)
        
        # 2. User Input (Round 1)
        doc.add_heading("2. User Input (Round 1)", level=2)
        
        if strategy == "stepwise":
            user_prompt = prompt_gen.get_step_prompt(1, mock_title, mock_abstract)
        else:
            user_prompt = prompt_gen.get_single_prompt(mock_title, mock_abstract)
            
        p = doc.add_paragraph(user_prompt)
        p.runs[0].font.name = 'Courier New'
        p.runs[0].font.size = Pt(10)
        p.runs[0].bold = True
        
        # 3. Simulated Interaction (for multi-step strategies)
        if strategy == "stepwise":
            doc.add_heading("3. Stepwise Interaction Chain", level=2)
            
            # Step 1 Response (Mock)
            doc.add_paragraph("Model Output (Step 1): 1", style="Quote")
            
            # Step 2 Input
            doc.add_paragraph("User Input (Step 2):")
            step2_prompt = prompt_gen.get_step_prompt(2, mock_title, mock_abstract)
            # Note: In real code, only the specific question is appended to history, 
            # but get_step_prompt returns full context if needed. 
            # Actually get_step_prompt in current code returns Full Title+Abstract+Question.
            # Let's show exactly what get_step_prompt returns.
            p = doc.add_paragraph(step2_prompt)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            
            # Step 2 Response (Mock)
            doc.add_paragraph("Model Output (Step 2): 1", style="Quote")
            
            # Step 3 Input
            doc.add_paragraph("User Input (Step 3):")
            step3_prompt = prompt_gen.get_step_prompt(3, mock_title, mock_abstract)
            p = doc.add_paragraph(step3_prompt)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(9)
            
        elif strategy == "reflection":
            doc.add_heading("3. Reflection Interaction Chain", level=2)
            
            # Round 1 Response (Mock)
            doc.add_paragraph("Model Output (Round 1): 1\t1\tCity\tChina-Shenzhen", style="Quote")
            
            # Round 2 Input (Critique)
            doc.add_paragraph("User Input (Round 2 - Critique):")
            critique_prompt = prompt_gen.get_reflection_critique_prompt()
            p = doc.add_paragraph(critique_prompt)
            p.runs[0].font.name = 'Courier New'
            p.runs[0].font.size = Pt(10)
            p.runs[0].bold = True

    # Save
    doc.save(output_file)
    print(f"Audit report saved to {output_file}")

if __name__ == "__main__":
    create_audit_report()
