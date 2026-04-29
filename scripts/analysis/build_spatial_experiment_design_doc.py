from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ANALYSIS_DIR = Path("Data/spatial_sample_2000_20260428/analysis")
DOCX_PATH = ANALYSIS_DIR / "spatial_extraction_experiment_design_20260428.docx"
MD_PATH = ANALYSIS_DIR / "spatial_extraction_experiment_design_20260428.md"


SECTIONS = [
    (
        "摘要",
        [
            "本文档系统阐述空间指标提取任务的实验设计、提取策略、后处理校验机制与有效性验证结果。该任务的目标并非泛化识别论文是否讨论空间议题，而是从题名与摘要中判断论文的主要研究是否在可识别的真实地理区域展开，并进一步抽取最大空间尺度与具体研究区域。为降低大模型在地理空间抽取中常见的过度推断和占位式幻觉，本方案采用“语义锁定 - 结构化输出 - 证据校验 - 审计回写”的四层架构：首先通过空间专用提示词约束研究区定义，其次要求模型输出严格 JSON，再通过本地解析器验证空间区域是否来自 Title/Abstract 或满足受限隐含国家/地区规则，最后在输出中保留 validation status、reason 和 area evidence。",
            "在 2000 篇随机样本文献上的端到端实验显示，pipeline 输出空间正例 1421 篇，非空间 579 篇；所有正例均满足 accepted 状态，未出现正例但研究区为 Not mentioned 的情况。进一步以 Codex GPT 盲标作为 pseudo-gold 进行独立诊断评估，有效盲标 2000/2000，pipeline 与 pseudo-gold 的二分类 Accuracy 为 0.9435，Precision 为 0.9852，Recall 为 0.9383，F1 为 0.9612，Cohen kappa 为 0.8575。结果表明，该方案在保持高召回的同时显著压制了未命名城市、泛称 site 和幻觉地名等高风险错误，具备作为当前空间指标提取稳定版本的可行性。",
        ],
    ),
    (
        "1. 研究问题与任务定义",
        [
            "空间指标提取任务服务于城市更新文献的计量分析，其核心问题是：一篇论文的实证研究、案例研究、数据采集或政策分析是否锚定在真实地理空间边界内；若是，其主要研究区域的最大空间尺度和具体区域名称是什么。该任务强调“研究展开区域”而不是“文本中出现过的所有地名”，因此必须排除理论背景、学派来源、比较参照、宏观语境和文献回顾中的地理噪声。",
            "本研究将空间研究定义为：论文的核心研究目标、经验分析、数据来源、案例对象或政策建议实质性锚定于地球上的真实物理/地理边界。相应地，纯理论推导、纯算法或工程模型、实验室研究、无特定地理边界的人群调查、以及只暗示存在案例地但未给出可识别区域的文本，均判定为非空间研究。",
            "该定义直接回应此前模型误判问题：摘要中出现 city、site、case study、brownfield development 等词，并不自动构成空间研究证据；只有当具体区域能够从 Title/Abstract 中命名或明确识别时，才允许输出空间正例。",
        ],
    ),
    (
        "2. 指标体系与输出合同",
        [
            "空间指标输出保持既有 Excel 合同，核心列包括 `空间研究/非空间研究`、`空间等级`、`具体空间描述`、`Reasoning` 和 `Confidence`。在此基础上，当前稳定版增加三个审计字段：`spatial_validation_status`、`spatial_validation_reason`、`spatial_area_evidence`。这些字段不改变核心统计口径，但为每条记录提供可追溯的保留或拒绝原因。",
            "空间尺度采用最大包络原则（Maximum Bounding Box Principle），即以论文真实研究区域的最大空间范围映射到 9 级尺度。若研究同时涉及多个城市、区域或国家，尺度按覆盖范围而非单个样例中最小地理单元判定。",
        ],
    ),
    (
        "3. 提取策略设计",
        [
            "本方案采用大模型语义判断与本地确定性校验相结合的策略。大模型负责处理题名和摘要中的复杂语义、区分研究对象和背景噪声；本地程序负责执行不可妥协的合同约束，包括 JSON 解析、布尔归一化、尺度归一化、泛称区域拦截、证据回源校验和尺度-区域一致性检查。",
            "提示词层面保留原有 Role、Core Guidelines、Scale Mapping Rules 和 JSON 输出格式，只在关键边界处补充最小约束：研究边界必须在 Title/Abstract 中命名或可识别；泛称 city/site/context 不足以构成空间研究；不得输出 “unspecified city”、“implicit city”、“unknown site”、“case study context” 等占位区域；隐含地理只能在明确机构、政策或国家/地区制度语境支持时用于国家/地区层面，不能推断未命名城市、场地、街区或项目区。",
            "解析层面将 `Is_Spatial_Research` 归一化为真实布尔值，并将非空间记录统一输出为 `0 / Not mentioned / Not mentioned`。当模型返回空间正例时，本地校验不会直接信任其地名，而是逐级验证区域值、尺度值和证据来源。",
        ],
    ),
    (
        "4. 后处理与证据校验机制",
        [
            "后处理是本方案可靠性的核心。空间提取中最危险的错误并不是模型完全无法识别空间，而是模型在文本没有明确区域时生成看似合理的占位地点或命名地点。因此，本方案将“是否空间研究”的最终落盘决策从纯模型输出改为模型输出与本地证据校验共同决定。",
            "第一，泛称区域过滤会拒绝空值、Not mentioned、unknown、unspecified、unnamed、not specified、case study context 等占位表达，并进一步识别 “A brownfield site”、“the study area in a city”、“the municipality under study” 等没有命名锚点的泛称边界。该策略避免模型把普通空间名词误当作研究区。",
            "第二，证据回源校验要求 `Specific_Study_Area` 必须能够在 Title/Abstract 中找到明确支持。系统依次检查完整短语命中、去括号/去隐含后缀的 primary anchor 命中、多个区域片段共同命中，以及受限隐含国家/地区证据。若模型给出一个看起来像地名但 Title/Abstract 中没有支撑，最终会被标记为 rejected，并回写 `area_not_supported_by_title_or_abstract`。",
            "第三，尺度-区域一致性校验用于拦截结构性矛盾。例如，United Kingdom、China 这类国家/地区不能与 Single-city / Municipal Scale 搭配；global/worldwide 只能映射到 Global Scale；含 implicit 的低层级城市、街区或项目区推断会被拒绝。",
            "第四，审计列将每条样本的最终状态显式暴露：accepted 表示空间正例且证据通过；not_spatial 表示模型和校验均落入非空间；rejected 表示模型尝试输出空间正例，但本地证据或尺度校验未通过。这样既保留了高质量正例，又为后续人工复核提供可定位样本。",
        ],
    ),
    (
        "5. 实验数据与运行设置",
        [
            "本轮稳定性验证采用 `Data/spatial_sample_2000_20260428/input/spatial_sample_2000_seed20260428.xlsx` 作为输入。样本从已清洗的 10000 篇测试集抽取，固定随机种子为 20260428，保留 Article Title 与 Abstract 作为核心证据来源。输入文件共 2000 行，题名和摘要均无空值，Title + Abstract 无完全重复。",
            "pipeline 主运行使用 `spatial-zero` 模板，输出文件为 `Data/spatial_sample_2000_20260428/analysis/spatial_zero_2000_final_v2_20260428.xlsx`。为检验策略是否达到稳定版要求，本轮额外调用 Codex GPT 对同一 2000 篇文献进行盲标。盲标只读取 Title 与 Abstract，不读取 pipeline 预测结果，输出 `gpt_is_spatial`、`gpt_spatial_scale_level` 和 `gpt_specific_study_area` 三个字段。该盲标作为 pseudo-gold 诊断基准，不等同于真实人工标签。",
        ],
    ),
    (
        "6. 实验结果",
        [
            "pipeline 最终输出中，空间正例为 1421 篇，占 71.05%；非空间为 579 篇，占 28.95%。空间验证状态中，accepted 为 1421，not_spatial 为 562，rejected 为 17。所有空间正例均对应 accepted 状态，且不存在正例研究区为 Not mentioned 的记录。",
            "validation_reason 分布显示，1294 条正例由完整 explicit_area_evidence 支撑，110 条由 explicit_area_fragment_evidence 支撑，17 条由 explicit_area_primary_anchor_evidence 支撑；562 条为 model_non_spatial；17 条被拒绝为 area_not_supported_by_title_or_abstract。该分布说明当前系统的正例主要依赖可回源的明确区域证据，而不是模型自由推断。",
            "Codex GPT pseudo-gold 盲标中，空间正例为 1492 篇，非空间为 508 篇，有效标签 2000/2000。与 pipeline 对比后，二分类混淆矩阵为 TP=1400、TN=487、FP=21、FN=92。对应 Accuracy=0.9435，Precision=0.9852，Recall=0.9383，Specificity=0.9587，F1=0.9612，Cohen kappa=0.8575。",
            "在双方均为空间正例的 1400 条样本中，空间尺度 exact match 为 1082 条，准确率为 0.7729。具体区域自动匹配中，exact 为 691，containment 为 389，token_overlap 为 251，different 为 69。二分类结果明显优于尺度与区域字符串完全一致性，说明当前策略已经较好解决“是否空间研究”的主任务，但尺度边界和区域字符串表达仍是后续精细化优化方向。",
        ],
    ),
    (
        "7. 可行性论证",
        [
            "第一，任务定义具备可操作性。空间研究被限定为核心研究行为与真实地理边界之间的实质性锚定关系，而不是普通地名出现。这一界定能被 Title/Abstract 证据支持，也能被本地规则进行后验校验。",
            "第二，技术路径具有稳定性。模型输出被限制为严格 JSON，本地解析器将所有字段归一化为固定 Excel 合同；即使模型返回字符串布尔值、文本尺度标签或附加解释，本地层仍可恢复为统一结构。对于无法解析、缺少尺度、泛称区域或证据不支持的样本，系统统一回落到非空间输出，并保留 rejection reason。",
            "第三，证据回源机制降低了大模型幻觉风险。不同于直接信任 LLM 的地名输出，本方案要求具体研究区域必须来自 Title/Abstract 或满足明确机构/政策支撑的受限隐含国家/地区规则。该机制直接拦截命名地点幻觉和未命名 case/site/city 推断。",
            "第四，实验结果支持方案可行。2000 篇样本中 pipeline 与 Codex GPT pseudo-gold 的 F1 达到 0.9612，Cohen kappa 达到 0.8575，说明系统输出与独立盲标之间具有较强一致性；同时 FP 仅 21 条，表明新增约束并未造成大量伪空间正例。",
        ],
    ),
    (
        "8. 优越性分析",
        [
            "相较于纯提示词方案，本方案的优势在于将语义理解与确定性验证分离。LLM 负责识别复杂语义关系，本地程序负责执行不可让步的边界条件。这样既保留了大模型对学术摘要复杂表达的适应性，又避免模型将隐含或泛称对象自由补全为具体地理区域。",
            "相较于纯规则或词典方案，本方案不会简单依赖地名词典或 city/site 等关键词。许多空间研究的区域表达具有复杂结构，例如 “40 large cities in France and Germany”、“Lisbon metropolitan area”、“seam line neighbourhood in Jerusalem, East Musrara”。LLM 能较好识别这些短语与研究行为的关系，而本地层只在其输出后进行证据与尺度约束。",
            "相较于早期未加证据校验的空间抽取，本方案显著提升了可解释性和可审计性。每条记录不仅有 1/0、尺度和研究区，还包含 `spatial_validation_status`、`spatial_validation_reason` 和 `spatial_area_evidence`。这使得错误分析可以定位到具体机制，例如 model_non_spatial、explicit_area_fragment_evidence 或 area_not_supported_by_title_or_abstract，而不是只能查看模型 Reasoning。",
            "从科研复现角度看，本方案具备稳定输入、固定输出合同、可复跑的样本文件、可追踪的 correction 记录和独立 pseudo-gold 评估。它不仅适合当前 2000 篇测试，也可扩展到后续 10000 篇或全量文献处理。",
        ],
    ),
    (
        "9. 局限性与后续改进",
        [
            "第一，Codex GPT 盲标是 pseudo-gold，而非真实人工标签。因此评估结果反映 pipeline 与另一个高能力模型的诊断一致性，不能直接等同于人工准确率。若用于正式论文中的监督性能表述，仍建议抽取冲突样本进行人工复核。",
            "第二，尺度 exact match 为 0.7729，低于二分类 F1。这表明空间尺度映射仍存在边界歧义，例如都市圈与单城、区县与微观街区、多城市与跨国比较之间可能存在不同标注口径。后续可以针对 Scale_Disagreements 建立人工 adjudication 规则。",
            "第三，具体区域字符串存在表达差异。containment 和 token_overlap 大量存在，说明模型与 pipeline 可能都识别了正确区域，但一个输出更长的地理短语，另一个输出核心地名。后续可以增加标准化地名层或 canonical area 字段，但不建议替代原文抽取字段，因为原文短语更利于审计。",
            "第四，隐含国家/地区规则当前保持保守，仅覆盖少量高置信别名和政策/机构上下文。这降低了误报风险，但也可能牺牲部分召回。若后续任务更重视召回，可在不放开未命名城市/site 的前提下扩展国家/地区机构词表。",
        ],
    ),
    (
        "10. 结论",
        [
            "当前空间指标提取方案已经形成可用于稳定生产的闭环：以 Title/Abstract 为唯一证据源，以 LLM 进行语义判断，以本地程序进行强约束校验，以审计列支撑错误追踪，并通过 2000 篇样本和 Codex GPT pseudo-gold 盲标完成独立诊断评估。",
            "实验结果显示，方案在二分类层面具有较高一致性和较低误报风险，能够有效解决此前“摘要没有明确地理空间描述却输出 unspecified city”的核心问题。考虑到输出合同稳定、证据校验充分、评估指标良好、错误样本可追溯，本文认为该空间指标提取策略可以作为当前项目的稳定版本，用于后续更大规模文献空间研究指标抽取。",
        ],
    ),
]

SCALE_ROWS = [
    ("1", "Global Scale", "全球范围研究或明确的 worldwide/global 分析"),
    ("2", "Multi-national / Continental Scale", "跨国、洲际或多个国家比较"),
    ("3", "National / Single-country Scale", "单一国家范围"),
    ("4", "Multi-provincial / Sub-national Regional Scale", "国家内部多个省/区域/流域/大区"),
    ("5", "Single-provincial / State Scale", "单一省、州或同级行政区"),
    ("6", "Multi-city / Megaregion Scale", "多个城市或都市圈/城市群"),
    ("7", "Single-city / Municipal Scale", "单一城市或市域"),
    ("8", "District / County Scale", "区县或同级行政区域"),
    ("9", "Micro / Neighborhood / Block Scale", "街区、社区、站点、地块、走廊等微观空间"),
]

METRICS_ROWS = [
    ("Pipeline spatial positive", "1421"),
    ("Pipeline spatial negative", "579"),
    ("Accepted / Not spatial / Rejected", "1421 / 562 / 17"),
    ("Codex GPT pseudo-gold valid labels", "2000 / 2000"),
    ("GPT pseudo-gold positive / negative", "1492 / 508"),
    ("TP / TN / FP / FN", "1400 / 487 / 21 / 92"),
    ("Accuracy", "0.9435"),
    ("Precision", "0.9852"),
    ("Recall", "0.9383"),
    ("Specificity", "0.9587"),
    ("F1", "0.9612"),
    ("Cohen kappa", "0.8575"),
    ("Scale exact match among mutual positives", "1082 / 1400 = 0.7729"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def write_markdown() -> None:
    lines = [
        "# 空间指标提取任务实验设计与方法论说明",
        "",
        "**面向城市更新文献的研究区域识别、空间尺度判定与可追溯验证**",
        "",
        "日期：2026-04-28",
        "",
    ]
    for title, paragraphs in SECTIONS:
        lines.append(f"## {title}")
        lines.append("")
        for paragraph in paragraphs:
            lines.append(paragraph)
            lines.append("")
    lines.extend(
        [
            "## 附表 1：空间尺度映射",
            "",
            "| 等级 | 标签 | 判定说明 |",
            "|---|---|---|",
        ]
    )
    for row in SCALE_ROWS:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} |")
    lines.extend(
        [
            "",
            "## 附表 2：关键实验指标",
            "",
            "| 指标 | 数值 |",
            "|---|---|",
        ]
    )
    for row in METRICS_ROWS:
        lines.append(f"| {row[0]} | {row[1]} |")
    lines.append("")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def write_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_heading("空间指标提取任务实验设计与方法论说明", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("面向城市更新文献的研究区域识别、空间尺度判定与可追溯验证")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(12)
    date_paragraph = doc.add_paragraph("日期：2026-04-28")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for heading, paragraphs in SECTIONS:
        doc.add_heading(heading, level=1)
        for text in paragraphs:
            paragraph = doc.add_paragraph(text)
            paragraph.paragraph_format.first_line_indent = Pt(21)
            paragraph.paragraph_format.line_spacing = 1.15

    doc.add_heading("附表 1：空间尺度映射", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(["等级", "标签", "判定说明"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for level, label, desc in SCALE_ROWS:
        cells = table.add_row().cells
        set_cell_text(cells[0], level)
        set_cell_text(cells[1], label)
        set_cell_text(cells[2], desc)

    doc.add_heading("附表 2：关键实验指标", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(["指标", "数值"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for key, value in METRICS_ROWS:
        cells = table.add_row().cells
        set_cell_text(cells[0], key)
        set_cell_text(cells[1], value)

    doc.add_paragraph("注：Codex GPT 盲标为 pseudo-gold 诊断基准，不等同于真实人工标签。")
    doc.save(DOCX_PATH)


def main() -> int:
    ANALYSIS_DIR.mkdir(parents=True, exist_ok=True)
    write_markdown()
    write_docx()
    print(DOCX_PATH)
    print(MD_PATH)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
