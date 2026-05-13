from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any, Dict

import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parents[3]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from scripts.analysis.spatial.blind_label_common import (  # noqa: E402
    area_match_type,
    binary_metrics,
    normalize_scale,
    normalize_spatial_flag,
    normalize_text,
)
from src.runtime.config import Schema  # noqa: E402


DEFAULT_ANALYSIS_DIR = PROJECT_ROOT / "Data" / "spatial_sample_2000_20260428" / "analysis"
DEFAULT_LABELS = DEFAULT_ANALYSIS_DIR / "codex_gpt_blind_labels_2000_20260428.xlsx"
DEFAULT_PIPELINE = DEFAULT_ANALYSIS_DIR / "spatial_zero_2000_final_v2_20260428.xlsx"
DEFAULT_OUTPUT_XLSX = DEFAULT_ANALYSIS_DIR / "spatial_codex_gpt_vs_pipeline_eval_20260428.xlsx"
DEFAULT_OUTPUT_JSON = DEFAULT_ANALYSIS_DIR / "spatial_codex_gpt_vs_pipeline_eval_20260428.json"
DEFAULT_OUTPUT_DOCX = DEFAULT_ANALYSIS_DIR / "spatial_codex_gpt_vs_pipeline_report_20260428.docx"


def normalize_pipeline_flag(value: Any) -> int:
    return int(normalize_spatial_flag(value))


def _with_row_index(frame: pd.DataFrame, column_name: str) -> pd.DataFrame:
    if column_name in frame.columns:
        raise ValueError(f"Input already contains reserved column {column_name!r}")
    return frame.reset_index(names=column_name)


def _add_title_key(frame: pd.DataFrame) -> pd.DataFrame:
    if Schema.TITLE not in frame.columns:
        raise RuntimeError(f"Cannot align by title: missing {Schema.TITLE!r}")
    out = frame.copy()
    out["_title_key"] = out[Schema.TITLE].map(normalize_text)
    if out["_title_key"].duplicated().any():
        duplicates = out.loc[out["_title_key"].duplicated(keep=False), Schema.TITLE].head(10).tolist()
        raise RuntimeError(f"Cannot align by title: duplicate titles found, examples={duplicates}")
    return out


def _row_titles_match(labels: pd.DataFrame, pipeline: pd.DataFrame) -> bool:
    if Schema.TITLE not in labels.columns or Schema.TITLE not in pipeline.columns:
        return False
    if len(labels) != len(pipeline):
        return False
    left = labels.sort_values("row_index")[Schema.TITLE].map(normalize_text).reset_index(drop=True)
    right = pipeline.sort_values("pipeline_row_index")[Schema.TITLE].map(normalize_text).reset_index(drop=True)
    return bool(left.equals(right))


def load_and_align(labels_path: Path, pipeline_path: Path, align_on: str = "auto") -> pd.DataFrame:
    labels = pd.read_excel(labels_path, engine="openpyxl")
    pipeline = pd.read_excel(pipeline_path, engine="openpyxl")
    if "row_index" not in labels.columns:
        labels = _with_row_index(labels, "row_index")
    labels["row_index"] = labels["row_index"].astype(int)
    pipeline = _with_row_index(pipeline, "pipeline_row_index")

    if align_on not in {"auto", "row_index", "title"}:
        raise ValueError("--align-on must be one of: auto, row_index, title")

    resolved_align = align_on
    if align_on == "auto":
        resolved_align = "row_index" if _row_titles_match(labels, pipeline) else "title"

    if resolved_align == "row_index":
        merged = labels.merge(
            pipeline.rename(columns={"pipeline_row_index": "row_index"}),
            on="row_index",
            how="left",
            suffixes=("_gpt_label_file", "_pipeline"),
            validate="one_to_one",
        )
    else:
        labels_by_title = _add_title_key(labels)
        pipeline_by_title = _add_title_key(pipeline)
        merged = labels_by_title.merge(
            pipeline_by_title,
            on="_title_key",
            how="left",
            suffixes=("_gpt_label_file", "_pipeline"),
            validate="one_to_one",
        )
        merged = merged.drop(columns=["_title_key"])

    if len(merged) != len(labels):
        raise RuntimeError(f"Aligned row count mismatch: labels={len(labels)} merged={len(merged)}")
    if merged[Schema.IS_SPATIAL].isna().any():
        missing = int(merged[Schema.IS_SPATIAL].isna().sum())
        raise RuntimeError(f"Aligned rows contain missing pipeline predictions: missing={missing}")
    merged.attrs["alignment_mode"] = resolved_align
    return merged


def prepare_eval_frame(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["gpt_valid"] = out.get("label_status", "").astype(str).str.lower().eq("valid")
    out["gpt_is_spatial_norm"] = out["gpt_is_spatial"].map(lambda value: int(normalize_spatial_flag(value)))
    out["pipeline_is_spatial_norm"] = out[Schema.IS_SPATIAL].map(normalize_pipeline_flag)
    out["gpt_scale_norm"] = out["gpt_spatial_scale_level"].map(normalize_scale)
    out["pipeline_scale_norm"] = out[Schema.SPATIAL_LEVEL].map(normalize_scale)
    out["area_match_type"] = out.apply(
        lambda row: area_match_type(row.get("gpt_specific_study_area"), row.get(Schema.SPATIAL_DESC)),
        axis=1,
    )
    out["binary_bucket"] = out.apply(binary_bucket, axis=1)
    return out


def binary_bucket(row: pd.Series) -> str:
    if not bool(row.get("gpt_valid")):
        return "invalid_gpt_label"
    gold = int(row.get("gpt_is_spatial_norm", 0))
    pred = int(row.get("pipeline_is_spatial_norm", 0))
    if gold == 1 and pred == 1:
        return "TP"
    if gold == 0 and pred == 0:
        return "TN"
    if gold == 0 and pred == 1:
        return "FP"
    return "FN"


def to_records_table(mapping: Dict[str, Any]) -> pd.DataFrame:
    return pd.DataFrame([{"metric": key, "value": value} for key, value in mapping.items()])


def build_summary(eval_df: pd.DataFrame) -> Dict[str, Any]:
    valid = eval_df[eval_df["gpt_valid"]].copy()
    metrics = binary_metrics(valid["gpt_is_spatial_norm"], valid["pipeline_is_spatial_norm"])
    both_positive = valid[(valid["gpt_is_spatial_norm"] == 1) & (valid["pipeline_is_spatial_norm"] == 1)].copy()
    scale_n = len(both_positive)
    scale_exact = int((both_positive["gpt_scale_norm"] == both_positive["pipeline_scale_norm"]).sum()) if scale_n else 0
    area_counts = both_positive["area_match_type"].value_counts(dropna=False).to_dict()
    return {
        "label_source": "Codex GPT pseudo-gold blind labels",
        "alignment_mode": str(eval_df.attrs.get("alignment_mode", "")),
        "total_labels": int(len(eval_df)),
        "valid_labels": int(len(valid)),
        "invalid_or_failed_labels": int((~eval_df["gpt_valid"]).sum()),
        "pipeline_rows_aligned": int(eval_df[Schema.IS_SPATIAL].notna().sum()),
        "binary_metrics": metrics,
        "scale_both_positive_n": int(scale_n),
        "scale_exact_match_n": scale_exact,
        "scale_exact_match_accuracy": scale_exact / scale_n if scale_n else 0.0,
        "area_match_counts": {str(k): int(v) for k, v in area_counts.items()},
        "binary_bucket_counts": {
            str(k): int(v) for k, v in eval_df["binary_bucket"].value_counts(dropna=False).to_dict().items()
        },
    }


def write_workbook(eval_df: pd.DataFrame, summary: Dict[str, Any], output: Path) -> None:
    valid = eval_df[eval_df["gpt_valid"]].copy()
    both_positive = valid[(valid["gpt_is_spatial_norm"] == 1) & (valid["pipeline_is_spatial_norm"] == 1)].copy()
    binary_confusion = pd.crosstab(
        valid["gpt_is_spatial_norm"],
        valid["pipeline_is_spatial_norm"],
        rownames=["GPT pseudo-gold"],
        colnames=["Pipeline"],
        dropna=False,
    )
    scale_confusion = pd.crosstab(
        both_positive["gpt_scale_norm"],
        both_positive["pipeline_scale_norm"],
        rownames=["GPT scale"],
        colnames=["Pipeline scale"],
        dropna=False,
    )
    scale_metrics = {
        "both_positive_n": summary["scale_both_positive_n"],
        "exact_match_n": summary["scale_exact_match_n"],
        "exact_match_accuracy": summary["scale_exact_match_accuracy"],
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        to_records_table(
            {
                "evaluation_type": "GPT pseudo-gold evaluation",
                "total_labels": summary["total_labels"],
                "valid_labels": summary["valid_labels"],
                "invalid_or_failed_labels": summary["invalid_or_failed_labels"],
                "pipeline_rows_aligned": summary["pipeline_rows_aligned"],
            }
        ).to_excel(writer, sheet_name="Overview", index=False)
        to_records_table(summary["binary_metrics"]).to_excel(writer, sheet_name="Binary_Metrics", index=False)
        binary_confusion.to_excel(writer, sheet_name="Binary_Confusion")
        to_records_table(scale_metrics).to_excel(writer, sheet_name="Scale_Metrics", index=False)
        scale_confusion.to_excel(writer, sheet_name="Scale_Confusion")
        pd.DataFrame(
            [{"match_type": key, "count": value} for key, value in summary["area_match_counts"].items()]
        ).to_excel(writer, sheet_name="Area_Match", index=False)
        eval_df[eval_df["binary_bucket"] == "FP"].to_excel(writer, sheet_name="Pipeline_Pos_GPT_Neg", index=False)
        eval_df[eval_df["binary_bucket"] == "FN"].to_excel(writer, sheet_name="Pipeline_Neg_GPT_Pos", index=False)
        both_positive[both_positive["gpt_scale_norm"] != both_positive["pipeline_scale_norm"]].to_excel(
            writer, sheet_name="Scale_Disagreements", index=False
        )
        both_positive[both_positive["area_match_type"].isin(["different", "missing"])].to_excel(
            writer, sheet_name="Area_Disagreements", index=False
        )
        eval_df[~eval_df["gpt_valid"]].to_excel(writer, sheet_name="GPT_Label_Quality_Issues", index=False)


def write_json_summary(summary: Dict[str, Any], output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")


def write_docx(summary: Dict[str, Any], output: Path) -> None:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("python-docx is required to write the DOCX report") from error

    metrics = summary["binary_metrics"]
    document = Document()
    document.add_heading("空间指标 Codex GPT 盲标对比评估报告", level=1)
    document.add_paragraph(
        "本报告将空间指标提取 pipeline 结果与 Codex GPT 盲标结果进行对比。"
        "Codex GPT 盲标仅作为 pseudo-gold 诊断基准，不等同于真实人工标签。"
    )
    document.add_heading("总体情况", level=2)
    overview_labels = {
        "total_labels": "盲标总数",
        "valid_labels": "有效盲标数",
        "invalid_or_failed_labels": "无效或失败数",
        "pipeline_rows_aligned": "成功对齐 pipeline 行数",
    }
    for key, label in overview_labels.items():
        document.add_paragraph(f"{label}: {summary[key]}")
    document.add_heading("二分类指标", level=2)
    metric_labels = {
        "accuracy": "Accuracy",
        "precision": "Precision",
        "recall": "Recall",
        "specificity": "Specificity",
        "f1": "F1",
        "cohen_kappa": "Cohen kappa",
    }
    for key, label in metric_labels.items():
        value = metrics[key]
        document.add_paragraph(f"{label}: {value:.4f}")
    document.add_paragraph(
        f"混淆矩阵计数: TP={metrics['tp']}, TN={metrics['tn']}, FP={metrics['fp']}, FN={metrics['fn']}。"
    )
    document.add_heading("尺度与具体区域", level=2)
    document.add_paragraph(
        f"双方均为空间正例时的尺度 exact match: "
        f"{summary['scale_exact_match_n']}/{summary['scale_both_positive_n']} "
        f"({summary['scale_exact_match_accuracy']:.4f})。"
    )
    document.add_paragraph(f"具体区域自动匹配分布: {summary['area_match_counts']}")
    document.add_heading("结论提示", level=2)
    document.add_paragraph(
        "二分类一致性较高，但仍存在 pipeline 负例而 Codex GPT 判为空间正例的 FN 样本，"
        "应优先查看评估工作簿中的 Pipeline_Neg_GPT_Pos sheet。"
    )
    document.add_paragraph(
        "尺度 exact match 低于二分类一致性，说明空间尺度口径仍是主要误差来源之一；"
        "Area_Disagreements sheet 中的 different 样本适合做下一轮人工复核。"
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Evaluate pipeline spatial output against Codex GPT pseudo-gold labels.")
    parser.add_argument("--labels", type=Path, default=DEFAULT_LABELS)
    parser.add_argument("--pipeline", type=Path, default=DEFAULT_PIPELINE)
    parser.add_argument("--output-xlsx", type=Path, default=DEFAULT_OUTPUT_XLSX)
    parser.add_argument("--output-json", type=Path, default=DEFAULT_OUTPUT_JSON)
    parser.add_argument("--output-docx", type=Path, default=DEFAULT_OUTPUT_DOCX)
    parser.add_argument(
        "--align-on",
        choices=["auto", "row_index", "title"],
        default="auto",
        help="Alignment key for labels and pipeline output. auto uses row_index only when titles already match by row.",
    )
    return parser


def main() -> int:
    args = build_arg_parser().parse_args()
    merged = load_and_align(args.labels, args.pipeline, align_on=args.align_on)
    eval_df = prepare_eval_frame(merged)
    eval_df.attrs["alignment_mode"] = merged.attrs.get("alignment_mode", "")
    summary = build_summary(eval_df)
    write_workbook(eval_df, summary, args.output_xlsx)
    write_json_summary(summary, args.output_json)
    write_docx(summary, args.output_docx)
    print(f"[INFO] eval workbook saved to {args.output_xlsx}")
    print(f"[INFO] eval json saved to {args.output_json}")
    print(f"[INFO] eval docx saved to {args.output_docx}")
    print(f"[INFO] binary_metrics={summary['binary_metrics']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
