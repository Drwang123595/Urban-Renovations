from __future__ import annotations

import tempfile
from datetime import datetime
from pathlib import Path
from typing import Optional

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .review_workbook_analysis import (
    ANALYSIS_SHEET_NAME,
    CHART_LAYOUTS,
    COL_SPACE_DETAIL,
    COL_SPATIAL_LEVEL,
    COL_TOPIC_NAME_ZH,
    COL_YEAR,
    KPI_SHEET_NAME,
    LEVEL_SHARE_SHEET,
    LEVEL_TOPIC_COUNT_SHEET,
    NOT_MENTIONED,
    SPACE_COUNT_SHEET,
    TOPIC_SPACE_TOP_SHEET,
    YEAR_LEVEL_COUNT_SHEET,
    YEAR_SPATIAL_FLAG_SHEET,
    YEAR_TOPIC_COUNT_SHEET,
    YEAR_TOPIC_SHARE_SHEET,
    create_review_analysis_chart_images,
    generate_review_analysis_workbook,
)


REPORT_TITLE = "城市更新实验结果可视化分析报告"


def generate_review_experiment_report(
    workbook_path: Path,
    *,
    output_path: Optional[Path] = None,
    ensure_analysis: bool = True,
) -> Path:
    workbook_path = Path(workbook_path)
    if ensure_analysis:
        generate_review_analysis_workbook(
            workbook_path,
            append=True,
            replace_analysis_sheets=True,
        )

    if output_path is None:
        output_path = workbook_path.with_name(f"{workbook_path.stem}_experiment_report.docx")
    output_path = Path(output_path)

    tables = load_report_tables(workbook_path)
    insights = summarize_report_insights(tables)

    document = Document()
    configure_document_styles(document)
    build_report_document(document, workbook_path, tables, insights)

    with tempfile.TemporaryDirectory(prefix="review_report_figs_") as tmp_dir_name:
        chart_paths = create_review_analysis_chart_images(tables, Path(tmp_dir_name))
        append_chart_sections(document, chart_paths, insights)
        document.save(output_path)

    return output_path


def load_report_tables(workbook_path: Path) -> dict[str, pd.DataFrame]:
    tables = {}
    for sheet_name in [
        KPI_SHEET_NAME,
        ANALYSIS_SHEET_NAME,
        YEAR_TOPIC_COUNT_SHEET,
        YEAR_TOPIC_SHARE_SHEET,
        YEAR_LEVEL_COUNT_SHEET,
        YEAR_SPATIAL_FLAG_SHEET,
        LEVEL_SHARE_SHEET,
        LEVEL_TOPIC_COUNT_SHEET,
        SPACE_COUNT_SHEET,
        TOPIC_SPACE_TOP_SHEET,
    ]:
        tables[sheet_name] = pd.read_excel(workbook_path, sheet_name=sheet_name, engine="openpyxl")
    return tables


def summarize_report_insights(tables: dict[str, pd.DataFrame]) -> dict[str, object]:
    analysis_df = tables[ANALYSIS_SHEET_NAME]
    year_topic_count = tables[YEAR_TOPIC_COUNT_SHEET]
    year_topic_share = tables[YEAR_TOPIC_SHARE_SHEET]
    year_spatial = tables[YEAR_SPATIAL_FLAG_SHEET]
    level_share = tables[LEVEL_SHARE_SHEET]
    space_count = tables[SPACE_COUNT_SHEET]
    topic_space_top = tables[TOPIC_SPACE_TOP_SHEET]

    topic_totals = (
        analysis_df[COL_TOPIC_NAME_ZH]
        .fillna("")
        .astype(str)
        .str.strip()
        .value_counts()
        .rename_axis(COL_TOPIC_NAME_ZH)
        .reset_index(name="Count")
    )
    topic_totals["Share"] = topic_totals["Count"] / len(analysis_df) if len(analysis_df) else 0.0

    valid_year_count = year_topic_count[year_topic_count[COL_YEAR] != "Total"].copy()
    valid_year_count[COL_YEAR] = valid_year_count[COL_YEAR].astype(int)
    valid_year_share = year_topic_share[year_topic_share[COL_YEAR] != "Total"].copy()
    valid_year_share[COL_YEAR] = valid_year_share[COL_YEAR].astype(int)

    topic_columns = [column for column in valid_year_count.columns if column not in [COL_YEAR, "Total"]]
    early_share = valid_year_share[valid_year_share[COL_YEAR] <= 2005][topic_columns].mean()
    late_share = valid_year_share[valid_year_share[COL_YEAR] >= 2016][topic_columns].mean()
    delta_share = (late_share - early_share).sort_values(ascending=False)

    spatial_year = year_spatial[year_spatial[COL_YEAR] != "Total"].copy()
    spatial_year[COL_YEAR] = spatial_year[COL_YEAR].astype(int)

    top_levels = level_share.sort_values("Count", ascending=False).reset_index(drop=True)
    top_spaces = space_count.sort_values("Count", ascending=False).reset_index(drop=True)

    topic_examples = {}
    for topic_name in topic_totals[COL_TOPIC_NAME_ZH].head(3).tolist():
        topic_examples[topic_name] = (
            topic_space_top[topic_space_top[COL_TOPIC_NAME_ZH] == topic_name]
            .head(3)[[COL_SPACE_DETAIL, "Count", "Share_Within_Topic"]]
            .to_dict("records")
        )

    return {
        "sample_count": int(len(analysis_df)),
        "year_min": int(valid_year_count[COL_YEAR].min()),
        "year_max": int(valid_year_count[COL_YEAR].max()),
        "topic_count": int(topic_totals.shape[0]),
        "spatial_count": int(spatial_year["Spatial"].sum()),
        "non_spatial_count": int(spatial_year["Non-spatial"].sum()),
        "spatial_share_avg": float(spatial_year["Spatial_Share"].mean()),
        "spatial_share_min_year": spatial_year.loc[spatial_year["Spatial_Share"].idxmin(), COL_YEAR],
        "spatial_share_min_value": float(spatial_year["Spatial_Share"].min()),
        "spatial_share_max_year": spatial_year.loc[spatial_year["Spatial_Share"].idxmax(), COL_YEAR],
        "spatial_share_max_value": float(spatial_year["Spatial_Share"].max()),
        "top_topics": topic_totals.head(10).copy(),
        "early_top_topics": (early_share.sort_values(ascending=False).head(5) * 100).round(2),
        "late_top_topics": (late_share.sort_values(ascending=False).head(5) * 100).round(2),
        "delta_top_topics": (delta_share.head(5) * 100).round(2),
        "delta_bottom_topics": (delta_share.tail(5) * 100).round(2),
        "top_levels": top_levels.head(8).copy(),
        "top4_level_share_spatial": float(top_levels.head(4)["Share_Spatial_Samples"].sum()),
        "top_spaces": top_spaces.head(10).copy(),
        "top10_space_share_detailed": float(top_spaces.head(10)["Share_Detailed_Spatial"].sum()),
        "distinct_space_count": int(top_spaces.shape[0]),
        "topic_examples": topic_examples,
    }


def configure_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_report_document(
    document: Document,
    workbook_path: Path,
    tables: dict[str, pd.DataFrame],
    insights: dict[str, object],
) -> None:
    title = document.add_heading(REPORT_TITLE, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"数据来源：{workbook_path.name}\n").font.color.rgb = RGBColor(90, 90, 90)
    meta.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}").font.color.rgb = RGBColor(90, 90, 90)

    document.add_heading("一、研究概况", level=1)
    overview = document.add_paragraph()
    overview.add_run(
        f"本次实验基于 {insights['sample_count']} 篇文献结果，时间跨度为 {insights['year_min']} 至 {insights['year_max']} 年，"
        f"共识别 {insights['topic_count']} 个主题。"
    )
    overview.add_run(
        f"其中空间研究 {insights['spatial_count']} 篇，占全部样本约 {insights['spatial_count'] / insights['sample_count']:.1%}；"
        f"非空间研究 {insights['non_spatial_count']} 篇。"
    )

    document.add_heading("二、核心结论", level=1)
    add_bullet(
        document,
        f"主题结构上，{format_topic_leader_sentence(insights['top_topics'])}，说明研究重心明显集中在城市更新的社会影响、居住区更新与内城再生议题。",
    )
    add_bullet(
        document,
        f"时间变化上，后期阶段（2016 年及以后）“{insights['delta_top_topics'].index[0]}”占比相较早期提升 {insights['delta_top_topics'].iloc[0]:.2f} 个百分点，"
        f"“{insights['delta_top_topics'].index[1]}”和“{insights['delta_top_topics'].index[2]}”也同步抬升，显示议题从一般治理逐渐转向更新治理、社区排斥与项目实施过程。",
    )
    add_bullet(
        document,
        f"空间维度上，空间研究平均占比为 {insights['spatial_share_avg']:.1%}，最低年份 {int(insights['spatial_share_min_year'])} 年仍达到 {insights['spatial_share_min_value']:.1%}。"
        f"空间等级以单城市、市辖区和国家尺度为主，其中前四类尺度合计占空间研究样本 {insights['top4_level_share_spatial']:.1%}，说明研究主要集中在可落地治理单元。",
    )
    add_bullet(
        document,
        f"具体空间方面，共识别 {insights['distinct_space_count']} 个不同空间描述，前 10 个高频空间仅占具名空间样本的 {insights['top10_space_share_detailed']:.1%}，"
        "表明案例分布较为分散，呈现典型的多城市、跨国家碎片化研究格局。",
    )

    document.add_heading("三、关键指标摘要", level=1)
    append_dataframe_table(document, insights["top_topics"].head(10), title="表 1 主题频次 Top 10", percent_columns={"Share"})
    append_dataframe_table(
        document,
        insights["top_levels"][[
            COL_SPATIAL_LEVEL,
            "Count",
            "Share_All_Samples",
            "Share_Spatial_Samples",
        ]],
        title="表 2 空间等级结构",
        percent_columns={"Share_All_Samples", "Share_Spatial_Samples"},
    )
    append_dataframe_table(
        document,
        insights["top_spaces"][[
            COL_SPACE_DETAIL,
            "Count",
            "Share_Detailed_Spatial",
            "First_Year",
            "Last_Year",
        ]],
        title="表 3 高频具体空间 Top 10",
        percent_columns={"Share_Detailed_Spatial"},
    )

    document.add_heading("四、图表解读", level=1)


def append_chart_sections(document: Document, chart_paths: list[tuple[str, Path]], insights: dict[str, object]) -> None:
    captions = build_chart_captions(insights)
    for index, (title, image_path) in enumerate(chart_paths, start=1):
        document.add_heading(f"图 {index} {title}", level=2)
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(image_path), width=Cm(15.8))

        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(captions.get(title, "")).italic = True

        explanation = document.add_paragraph()
        explanation.add_run("说明：").bold = True
        explanation.add_run(captions.get(title, ""))

    document.add_heading("五、结论与建议", level=1)
    add_bullet(
        document,
        "后续若希望观察更细的议题迁移，可以继续将年份划分为更稳定的阶段窗口，例如 5 年期或政策节点期，再比较主题结构变化。",
    )
    add_bullet(
        document,
        "具体空间当前保留了原始表述，适合全文检索和人工核查；若下一步用于发表或展示，建议再做一轮地名归一化，提高跨城市比较的可读性。",
    )
    add_bullet(
        document,
        "空间等级分析显示研究重心明显偏向单城市与区县尺度，后续可单独抽取跨城市、多国样本，分析其主题是否更偏政策扩散或比较研究。",
    )


def build_chart_captions(insights: dict[str, object]) -> dict[str, str]:
    top_topics = insights["top_topics"]
    first_topic = top_topics.iloc[0]
    second_topic = top_topics.iloc[1]
    top_level = insights["top_levels"].iloc[0]
    top_space = insights["top_spaces"].iloc[0]

    return {
        "Year-Topic Share": (
            f"图中显示“{first_topic[COL_TOPIC_NAME_ZH]}”在各年份中持续保持主导，并在 2016 年后进一步走强；"
            f"“{second_topic[COL_TOPIC_NAME_ZH]}”位居第二梯队，后期占比也明显抬升。"
        ),
        "Year-Topic Count Heatmap": (
            "热力图用于识别主题出现的时间集中带。颜色越深表示该年份该主题样本越多，可直观看到哪些议题具有持续积累，哪些只在局部年份活跃。"
        ),
        "Spatial Level Share": (
            f"单城市尺度是最主要的研究层级，共 {int(top_level['Count'])} 篇，在空间研究子集中占 {top_level['Share_Spatial_Samples']:.1%}。"
            "图中同时给出占全部样本与占空间研究子样本两种口径，便于同时判断总体重要性和空间内部结构。"
        ),
        "Spatial Level vs Topic": (
            "该图用于观察不同空间等级与主题的耦合关系。通常单城市、区县和国家尺度聚集了最主要的城市更新主题，而更高层级尺度的议题分布相对离散。"
        ),
        "Top Detailed Spaces": (
            f"高频具体空间中，“{top_space[COL_SPACE_DETAIL]}”出现 {int(top_space['Count'])} 次，排名第一；"
            "但前 10 个空间仅覆盖少部分具名空间样本，说明案例总体较分散。"
        ),
        "Year vs Spatial Level": (
            "该图反映不同年份中研究尺度的分布变化。整体来看，单城市尺度长期占优，区县尺度和国家尺度构成第二层级，表明研究仍以可操作的治理空间为核心。"
        ),
    }


def append_dataframe_table(
    document: Document,
    frame: pd.DataFrame,
    *,
    title: str,
    percent_columns: Optional[set[str]] = None,
) -> None:
    percent_columns = percent_columns or set()
    document.add_paragraph(title).runs[0].bold = True
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for idx, column_name in enumerate(frame.columns):
        header_cells[idx].text = str(column_name)

    for _, row in frame.iterrows():
        cells = table.add_row().cells
        for idx, column_name in enumerate(frame.columns):
            value = row[column_name]
            if pd.isna(value):
                text = ""
            elif column_name in percent_columns and isinstance(value, (float, int)):
                text = f"{value:.1%}"
            else:
                text = str(value)
            cells[idx].text = text
    document.add_paragraph("")


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(text, style="List Bullet")


def format_topic_leader_sentence(top_topics: pd.DataFrame) -> str:
    first = top_topics.iloc[0]
    second = top_topics.iloc[1]
    third = top_topics.iloc[2]
    return (
        f"“{first[COL_TOPIC_NAME_ZH]}”是最核心主题，共 {int(first['Count'])} 篇，占比 {first['Share']:.1%}；"
        f"其后依次为“{second[COL_TOPIC_NAME_ZH]}”({int(second['Count'])} 篇，{second['Share']:.1%})"
        f"和“{third[COL_TOPIC_NAME_ZH]}”({int(third['Count'])} 篇，{third['Share']:.1%})"
    )
